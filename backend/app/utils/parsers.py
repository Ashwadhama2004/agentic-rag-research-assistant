"""
File Parsers
Parse various document formats into text.
"""
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import io


class BaseParser:
    """Base class for document parsers."""
    
    supported_extensions: List[str] = []
    
    def parse(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse file content into text.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        raise NotImplementedError
    
    def parse_pages(self, file_content: bytes, filename: str) -> Tuple[List[str], Dict[str, Any]]:
        """
        Parse file into separate pages.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
        
        Returns:
            Tuple of (list of page texts, metadata)
        """
        text, metadata = self.parse(file_content, filename)
        return [text], metadata


class PDFParser(BaseParser):
    """Parser for PDF documents."""
    
    supported_extensions = [".pdf"]
    
    def parse(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(file_content))
            
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            
            metadata = {
                "page_count": len(reader.pages),
                "title": reader.metadata.title if reader.metadata else None,
                "author": reader.metadata.author if reader.metadata else None
            }
            
            return "\n\n".join(pages), metadata
            
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install pypdf2")
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    def parse_pages(self, file_content: bytes, filename: str) -> Tuple[List[str], Dict[str, Any]]:
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(file_content))
            
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                pages.append(text if text else "")
            
            metadata = {
                "page_count": len(reader.pages),
                "title": reader.metadata.title if reader.metadata else None,
                "author": reader.metadata.author if reader.metadata else None
            }
            
            return pages, metadata
            
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install pypdf2")


class DOCXParser(BaseParser):
    """Parser for Word documents."""
    
    supported_extensions = [".docx", ".doc"]
    
    def parse(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_content))
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    paragraphs.append(" | ".join(cells))
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            return "\n\n".join(paragraphs), metadata
            
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")


class TextParser(BaseParser):
    """Parser for plain text files."""
    
    supported_extensions = [".txt", ".text"]
    
    def parse(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = file_content.decode(encoding)
                metadata = {
                    "encoding": encoding,
                    "char_count": len(text),
                    "line_count": text.count('\n') + 1
                }
                return text, metadata
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with supported encodings")


class MarkdownParser(BaseParser):
    """Parser for Markdown files."""
    
    supported_extensions = [".md", ".markdown"]
    
    def parse(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        text = file_content.decode('utf-8')
        
        # Extract headings for structure
        import re
        headings = re.findall(r'^#+\s+(.+)$', text, re.MULTILINE)
        
        metadata = {
            "char_count": len(text),
            "line_count": text.count('\n') + 1,
            "headings": headings[:10]  # First 10 headings
        }
        
        return text, metadata


class DocumentParser:
    """
    Main document parser that routes to appropriate parser based on file type.
    """
    
    def __init__(self):
        self.parsers: Dict[str, BaseParser] = {}
        self._register_default_parsers()
    
    def _register_default_parsers(self):
        """Register default parsers."""
        parsers = [PDFParser(), DOCXParser(), TextParser(), MarkdownParser()]
        
        for parser in parsers:
            for ext in parser.supported_extensions:
                self.parsers[ext.lower()] = parser
    
    def register(self, extension: str, parser: BaseParser):
        """Register a custom parser for an extension."""
        self.parsers[extension.lower()] = parser
    
    def get_parser(self, filename: str) -> Optional[BaseParser]:
        """Get the appropriate parser for a file."""
        ext = Path(filename).suffix.lower()
        return self.parsers.get(ext)
    
    def parse(
        self,
        file_content: bytes,
        filename: str
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a file into text.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        parser = self.get_parser(filename)
        
        if not parser:
            ext = Path(filename).suffix
            raise ValueError(f"Unsupported file type: {ext}")
        
        text, metadata = parser.parse(file_content, filename)
        metadata["filename"] = filename
        metadata["file_type"] = Path(filename).suffix.lower()
        
        return text, metadata
    
    def parse_pages(
        self,
        file_content: bytes,
        filename: str
    ) -> Tuple[List[str], Dict[str, Any]]:
        """
        Parse a file into separate pages.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
        
        Returns:
            Tuple of (list of page texts, metadata)
        """
        parser = self.get_parser(filename)
        
        if not parser:
            ext = Path(filename).suffix
            raise ValueError(f"Unsupported file type: {ext}")
        
        pages, metadata = parser.parse_pages(file_content, filename)
        metadata["filename"] = filename
        metadata["file_type"] = Path(filename).suffix.lower()
        
        return pages, metadata
    
    @property
    def supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.parsers.keys())


# Global parser instance
_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """Get the global document parser instance."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
